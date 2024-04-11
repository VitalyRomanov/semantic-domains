import json
from pathlib import Path
from typing import Union

from docx import Document

from definitions import Domain, Question, Word


def read_docx(path):
    document = Document(path)
    return document


class RWCParser:
    def __init__(self, document_path: Union[str, Path]):
        self.document = read_docx(document_path)
        self.cursor = 0

    def document_len(self):
        return 100  # len(self.document.paragraphs)
    
    def get_current_line(self):
        if self.cursor < self.document_len():
            if self.cursor % 1 == 0:
                print(f"Parsing progress {self.cursor}/{self.document_len()}\r", end="")
            line = self.document.paragraphs[self.cursor]
            while line.text.strip() == "":
                self.advance()
                line = self.document.paragraphs[self.cursor]
            return line
        else:
            raise StopIteration()
    
    def advance(self):
        self.cursor += 1
    
    def parse_next_domain_title(self):
        line = self.get_current_line()
        assert line.style.name == "Heading 2"
        domain_code, *domain_title_parts = line.text.split(" ")
        domain_title = " ".join(domain_title_parts)
        self.advance()
        return domain_code, domain_title
    
    def parse_next_domain_description(self):
        line = self.get_current_line()
        assert line.style.name == "descr"
        domain_description = line.text
        self.advance()
        return domain_description
    
    def parse_next_question_text(self):
        line = self.get_current_line()
        assert line.style.name == "quest1"
        question_text = line.text
        question_num, *question_parts = question_text.split(" ")
        assert question_num.startswith("(") and question_num.endswith(")")
        question_num = question_num.lstrip("(").rstrip(")")
        question_text = " ".join(question_parts)
        self.advance()
        return int(question_num), question_text
    
    def parse_next_set_of_words(self):
        line = self.get_current_line()
        assert line.style.name == "words"
        words_text = line.text
        assert words_text.startswith("•")
        words_text = words_text.lstrip("•").strip()

        char_protection = {
            " ": "_",
            ",": "#"
        }

        def protect_brackets(string):
            stack = []
            modified_string = list(string)
            
            opening = {"("}
            correspondence = {")": "("}
            for char_ind in range(len(string)):
                char = string[char_ind]
                if char in opening:
                    stack.append(char)
                elif char in correspondence:
                    pair = stack.pop(-1)
                    assert pair == correspondence[char]
                else:
                    if len(stack) > 0 and char in char_protection:
                        modified_string[char_ind] = char_protection[char]
            return "".join(modified_string)
            
        words = list(
            filter(
                lambda x: x != "", 
                map(
                    lambda word: word.strip(), 
                    protect_brackets(words_text).split(",")
                )
            )
        )

        final_words = []
        for word in words:
            for c, p in char_protection.items():
                word = word.replace(p, c)
            final_words.append(Word(word, source="rwc"))
        
        self.advance()
        return final_words
    
    def parse_next_question(self):
        question_num, question_text = self.parse_next_question_text()
        question_words = self.parse_next_set_of_words()
        return Question(question_num, question_text, question_words, source="rwc")
    
    def parse_domain_questions(self):
        questions = []

        line = self.get_current_line()
        while line.style.name in ["quest1", "words"]:
            questions.append(self.parse_next_question())
            line = self.get_current_line()

        return questions

    def parse_next_domain(self):
        domain_code, domain_title = self.parse_next_domain_title()
        domain_description = self.parse_next_domain_description()
        questions = self.parse_domain_questions()

        return Domain(domain_code, domain_title, domain_description, questions, source="rwc")

    def parse(self):
        domains = []

        self.cursor = 0
        while self.cursor < self.document_len():
            try:
                domain = self.parse_next_domain()
            except StopIteration:
                break
            domains.append(domain)

        return domains


def parse_rwc_domains(domains_path: Union[str, Path]):
    parser = RWCParser(domains_path)
    return parser.parse()


def convert_rwc_domains_to_json(
        domains_path: Union[str, Path], output_json_path: Union[str, Path]
    ):
    domains = parse_rwc_domains(domains_path)

    dict_domains = []
    for domain in domains:
        dict_domains.append(domain.to_dict())

    with open(output_json_path, "w") as f:
        f.write(json.dumps(
            {
                "version": "0.1",
                "domains": dict_domains
            }, 
            indent=4
        ))
    