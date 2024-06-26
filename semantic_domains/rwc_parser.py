import json
from pathlib import Path
from typing import Dict, List, Tuple, Union
from importlib.metadata import version

from docx import Document

from semantic_domains.definitions import Domain, Question


def read_docx(path: Union[str, Path]):
    document = Document(str(path))
    return document


def reverse_dict(dict_: Dict) -> Dict:
    keys, vals = zip(*dict_.items())
    reversed = dict(zip(vals, keys))
    assert len(dict_) == len(reversed)
    return reversed


class RWCParser:
    """
    Parses a RWC domain docx  document into a list of domains.
    """
    domain_header_style_name = "Heading 2"
    domain_description_style_name = "descr"
    question_style_name = "quest1"
    word_style_name = "words"
    
    def __init__(self, document_path: Union[str, Path]):
        """
        Args:
            document_path (Union[str, Path]): The path to the document to parse.
        """
        self.document = read_docx(document_path)
        self.cursor = 0

    def num_lines(self) -> int:
        """
        Returns:
            int: The number of lines in the document.
        """
        return len(self.document.paragraphs)
    
    def get_current_line(self, print_progress_every_n_lines:int = 1):
        num_lines = self.num_lines()
        if self.cursor < num_lines:
            if self.cursor % print_progress_every_n_lines == 0:
                print(f"Parsing progress {self.cursor}/{num_lines}\r", end="")
            line = self.document.paragraphs[self.cursor]
            while line.text.strip() == "":  # skip empty lines
                # normally there are no empty lines, but this could still happen
                self.advance()
                line = self.document.paragraphs[self.cursor]
            return line
        else:
            raise StopIteration()
    
    def advance(self) -> None:
        self.cursor += 1

    def verify_style(self, line, style_name) -> None:
        assert line.style.name == style_name, f"Domain title style is expected to be `{style_name}`, but the current line has style `{line.style.name}`"
    
    def parse_domain_title(self) -> Tuple[str, str]:
        """
        Domain title in the document has `Heading 2` style, so here it is checked that the current line contains text in this style.
        """
        line = self.get_current_line()
        self.verify_style(line, self.domain_header_style_name)
        domain_code, *domain_title_parts = line.text.split(" ")
        domain_title = " ".join(domain_title_parts)  # reassemble the domain title
        self.advance()
        return domain_code, domain_title
    
    def parse_domain_description(self) -> str:
        """
        Domain description in the document has `descr` style, so here it is checked that the current line contains text in this style.

        Returns:
            str: The domain description.
        """
        line = self.get_current_line()
        self.verify_style(line, self.domain_description_style_name)
        domain_description = line.text
        self.advance()
        return domain_description
    
    def parse_question_text(self) -> Tuple[int, str]:
        line = self.get_current_line()
        self.verify_style(line, self.question_style_name)
        question_text = line.text
        question_num, *question_parts = question_text.split(" ")
        assert question_num.startswith("(") and question_num.endswith(")")
        question_num = question_num.lstrip("(").rstrip(")")
        question_text = " ".join(question_parts).strip()
        self.advance()
        return int(question_num), question_text
    
    def parse_set_of_words(self) -> List[str]:
        """
        Parses the list of words for the current question. The line that contains the words has `words` style, 
        so here it is checked that the current line contains text in this style. The line with words should start with `•`, 
        an error is raised if this is not the case. Words are separated by commas, but sometimes a word has some 
        associated content in brackets. The text in brackets can also contain commas. To make sure that the content inside the brackets 
        remains intact, the function `protect_brackets` is used. The commas are returned back after the line is split into words.
        Sometimes words have associated `(v)` marker, that signifies that the word is a verb. This marker is replaced with `(verb)`.

        Returns:
            List[str]: The list of words.
        """
        line = self.get_current_line()
        self.verify_style(line, self.word_style_name)
        words_text = line.text
        assert words_text.startswith("•")  # line that contains words starts with `•`
        words_text = words_text.lstrip("•").strip()

        # need to properly tokenize words (they are separated by commas), make sure to keep the content inside the brackets intact
        char_protection_replacement = {
            " ": "_",
            ",": "#"
        }

        def protect_brackets(string, replacement_map):
            """
            Protect characters used for splitting a string by replacing them with something else.
            """
            stack = []
            modified_string = list(string)
            
            opening = {"("}
            correspondence = {")": "("}  # this is the only type of brackets in the document
            for char_ind in range(len(string)):
                char = string[char_ind]
                if char in opening:
                    stack.append(char)
                elif char in correspondence:
                    pair = stack.pop(-1)
                    assert pair == correspondence[char]
                else:
                    if len(stack) > 0 and char in replacement_map:
                        modified_string[char_ind] = replacement_map[char]
            return "".join(modified_string)
        
        def recover_replaced_characters(word, replacement_map):
            for k, v in replacement_map.items():
                word = word.replace(v, k)
            return word
        
        def tokenize_words(words_text, replacement_map):
            words_text = protect_brackets(words_text, replacement_map)
            words = words_text.split(",")  # TODO INTJ can be split using exclamation mark
            words = [word.strip() for word in words]
            words = [word for word in words if word!= ""]  # remove empty words
            words = [word.replace("(v)", "(verb)") for word in words]  # replace (v) with (verb)
            words = [word.replace("(n)", "(noun)") for word in words]  # replace (n) with (noun)
            words = [recover_replaced_characters(word, replacement_map) for word in words]        
            return words

        final_words = tokenize_words(words_text=words_text, replacement_map=char_protection_replacement)
        
        self.advance()
        return final_words
    
    def parse_question(self):
        """
        Parses the current question. Question is represented by the question number and the 
        question text.

        Returns:
            Tuple[int, str]: The question number and the question text.
        """
        question_num, question_text = self.parse_question_text()
        question_words = self.parse_set_of_words()
        return Question(question_num, question_text, question_words)
    
    def parse_domain_questions(self) -> List[Question]:
        """
        Parse a block of questions for a domain. The should contain only questions and words with 
        corresponding text styles. 

        Returns:
            List[Question]: The list of questions.
        """
        questions = []

        line = self.get_current_line()
        while line.style.name in [self.question_style_name, self.word_style_name]:
            questions.append(self.parse_question())
            try:
                line = self.get_current_line()
            except StopIteration:
                break  # end of document

        return questions

    def parse_domain(self):
        """
        Parses a domain block. The domain block is expected to have a domain title, a domain description, 
        and a list of questions. If this order is not observed, an error is raised.

        Returns:
            Domain: The parsed domain.
        """
        domain_code, domain_title = self.parse_domain_title()
        domain_description = self.parse_domain_description()
        questions = self.parse_domain_questions()

        return Domain(domain_code, domain_title, domain_description, questions)

    def parse(self) -> List[Domain]:
        """
        Parses the document and returns a list of domains.

        Returns:
            List[Domain]: The list of domains.
        """
        domains = []

        self.cursor = 0
        while self.cursor < self.num_lines():
            domain = self.parse_domain()
            domains.append(domain)

        return domains


def parse_rwc_domains(domains_path: Union[str, Path]) -> List[Domain]:
    parser = RWCParser(domains_path)
    return parser.parse()


def dump_domains_to_json(domains: List[Domain], output_json_path: Union[str, Path]) -> None:
    with open(output_json_path, "w") as f:
        f.write(json.dumps(
            {
                "version": version("semantic_domains"),
                "domains": [domain.to_dict() for domain in domains]
            }, 
            indent=4
        ))
    

def convert_rwc_domains_to_json(
        domains_path: Union[str, Path], output_json_path: Union[str, Path]
    ) -> None:
    domains = parse_rwc_domains(domains_path)

    dict_domains = []
    for domain in domains:
        dict_domains.append(domain.to_dict())

    dump_domains_to_json(domains=dict_domains, output_json_path=output_json_path)

    # TODO need to do a verification by checking the average word length and then checking that all the words do not appear as outliers
